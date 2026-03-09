"""
Text extraction helpers.
"""

import logging
import uuid
import zipfile
from pathlib import Path

import pdfplumber
from docx import Document

from backend.file_utils import OUTPUT_DIR, sanitize_filename

logger = logging.getLogger(__name__)
_RAPID_OCR_ENGINE = None


def extract_text(input_path: str) -> tuple[str, str]:
    """
    Extract text from supported files and save it as a UTF-8 .txt file.
    """
    source = Path(input_path)
    ext = source.suffix.lower()
    safe_stem = Path(sanitize_filename(source.name)).stem or "extracted"
    output_path = OUTPUT_DIR / f"{safe_stem}_extracted_{uuid.uuid4().hex[:8]}.txt"

    logger.info("Extracting text from %s", source)

    if ext == ".pdf":
        text = _extract_pdf_text(source)
    elif ext in {".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tiff", ".webp"}:
        text = _extract_image_text(source)
    elif ext == ".docx":
        text = _extract_docx_text(source)
    elif ext == ".txt":
        text = source.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file format for text extraction: {ext}")

    if not text.strip():
        if ext == ".pdf":
            raise RuntimeError(
                "No embedded text was found in this PDF. If it is a scanned PDF, install the Tesseract OCR engine to extract text from images."
            )
        raise ValueError("No readable text was found in the uploaded file")

    output_path.write_text(text, encoding="utf-8")
    logger.info("Text saved: %s (%s words)", output_path, len(text.split()))
    return text, str(output_path)


def _extract_pdf_text(source: Path) -> str:
    page_texts: list[str] = []
    ocr_pages: list[int] = []

    with pdfplumber.open(source) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text()
            if page_text:
                page_texts.append(f"--- Page {index} ---\n{page_text}\n")
            else:
                ocr_pages.append(index)

    if ocr_pages:
        try:
            ocr_results = _extract_pdf_text_with_ocr(source, ocr_pages)
        except RuntimeError as exc:
            if not page_texts:
                raise
            logger.warning("OCR fallback skipped for %s: %s", source.name, exc)
        else:
            for page_number in ocr_pages:
                page_text = ocr_results.get(page_number, "").strip()
                if page_text:
                    page_texts.append(f"--- Page {page_number} ---\n{page_text}\n")

    return "\n".join(page_texts).strip()


def _extract_image_text(source: Path) -> str:
    from PIL import Image
    image = Image.open(source)
    return _ocr_image(image)


def _extract_docx_text(source: Path) -> str:
    try:
        document = Document(source)
    except zipfile.BadZipFile as exc:
        raise ValueError("Invalid DOCX file") from exc

    chunks: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                chunks.append(row_text)

    return "\n".join(chunks)


def _extract_pdf_text_with_ocr(source: Path, page_numbers: list[int]) -> dict[int, str]:
    images = _render_pdf_pages_for_ocr(source, page_numbers)
    results: dict[int, str] = {}
    for page_number, image in images.items():
        results[page_number] = _ocr_image(image)
    return results


def _render_pdf_pages_for_ocr(source: Path, page_numbers: list[int]) -> dict[int, object]:
    try:
        import pypdfium2 as pdfium
    except ImportError:
        return _render_pdf_pages_for_ocr_with_pdf2image(source, page_numbers)

    pdf = pdfium.PdfDocument(str(source))
    try:
        rendered: dict[int, object] = {}
        for page_number in page_numbers:
            page_index = page_number - 1
            page = pdf[page_index]
            try:
                bitmap = page.render(scale=2.0)
                rendered[page_number] = bitmap.to_pil()
            finally:
                page.close()
        return rendered
    finally:
        pdf.close()


def _render_pdf_pages_for_ocr_with_pdf2image(source: Path, page_numbers: list[int]) -> dict[int, object]:
    try:
        from pdf2image import convert_from_path
    except ImportError as exc:
        raise RuntimeError("PDF OCR requires pypdfium2 or pdf2image to render pages") from exc

    rendered: dict[int, object] = {}
    for page_number in page_numbers:
        try:
            images = convert_from_path(str(source), first_page=page_number, last_page=page_number)
        except Exception as exc:
            raise RuntimeError("PDF OCR rendering failed. Install Poppler if pdf2image is being used.") from exc
        if images:
            rendered[page_number] = images[0]
    return rendered


def _ensure_tesseract_available(pytesseract_module: object) -> None:
    try:
        pytesseract_module.get_tesseract_version()
    except Exception as exc:
        raise RuntimeError(
            "Tesseract OCR is not installed or not on PATH. Install the Tesseract engine to extract text from scanned PDFs and images."
        ) from exc


def _ocr_image(image: object) -> str:
    tesseract_error: Exception | None = None

    try:
        import pytesseract
    except ImportError:
        pytesseract = None

    if pytesseract is not None:
        try:
            _ensure_tesseract_available(pytesseract)
            return pytesseract.image_to_string(image)
        except Exception as exc:
            tesseract_error = exc

    try:
        from rapidocr_onnxruntime import RapidOCR
    except ImportError as exc:
        if tesseract_error:
            raise RuntimeError(
                "No OCR engine is available. Install Tesseract or rapidocr-onnxruntime to extract text from scanned PDFs and images."
            ) from tesseract_error
        raise RuntimeError(
            "No OCR engine is available. Install Tesseract or rapidocr-onnxruntime to extract text from scanned PDFs and images."
        ) from exc

    global _RAPID_OCR_ENGINE
    if _RAPID_OCR_ENGINE is None:
        _RAPID_OCR_ENGINE = RapidOCR()

    engine = _RAPID_OCR_ENGINE
    import numpy as np

    result, _ = engine(np.array(image))
    if not result:
        return ""
    return "\n".join(line[1] for line in result if len(line) > 1 and line[1]).strip()
